"""
Agent 1 — Document Extraction Agent
Responsibilities: Parse PDF/DOCX, extract text, clean, chunk for downstream agents.
"""
import re
import time
import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


@dataclass
class DocumentChunk:
    index: int
    text: str
    char_start: int
    char_end: int
    section_ref: Optional[str] = None     # e.g. "Section 3.2"


@dataclass
class ExtractionResult:
    raw_text: str
    chunks: list[DocumentChunk]
    page_count: int
    word_count: int
    char_count: int
    duration_ms: int
    file_type: str


CHUNK_SIZE = 800        # characters per chunk
CHUNK_OVERLAP = 150     # overlap for context continuity

# Regex to detect section headings like "1.", "3.2", "SECTION 4"
_SECTION_RE = re.compile(
    r"(?:^|\n)\s*(?:SECTION\s+)?(\d+(?:\.\d+)?)\s*[.:\-–]\s*(.+?)(?:\n|$)",
    re.IGNORECASE,
)


class DocumentExtractionAgent:
    """
    Extracts text from PDF and DOCX files, then chunks for the pipeline.
    Uses PyMuPDF (fitz) for PDFs and python-docx for Word documents.
    """

    name = "DocumentExtractionAgent"

    async def run(self, file_path: str) -> ExtractionResult:
        t0 = time.perf_counter()
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext == ".pdf":
            raw_text, page_count, file_type = self._extract_pdf(file_path)
        elif ext in (".docx", ".doc"):
            raw_text, page_count, file_type = self._extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

        cleaned = self._clean_text(raw_text)
        chunks = self._chunk_text(cleaned)

        duration_ms = round((time.perf_counter() - t0) * 1000)
        logger.info(f"[{self.name}] Extracted {len(cleaned)} chars, {len(chunks)} chunks in {duration_ms}ms")

        return ExtractionResult(
            raw_text=cleaned,
            chunks=chunks,
            page_count=page_count,
            word_count=len(cleaned.split()),
            char_count=len(cleaned),
            duration_ms=duration_ms,
            file_type=file_type,
        )

    # ── PDF ──────────────────────────────────────────────────

    def _extract_pdf(self, file_path: str) -> tuple[str, int, str]:
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            pages = []
            for page in doc:
                pages.append(page.get_text("text"))
            doc.close()
            return "\n\n".join(pages), len(pages), "pdf"
        except ImportError:
            logger.warning("PyMuPDF not installed, falling back to stub extraction.")
            return self._stub_text(), 1, "pdf"

    # ── DOCX ─────────────────────────────────────────────────

    def _extract_docx(self, file_path: str) -> tuple[str, int, str]:
        try:
            from docx import Document
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            return "\n\n".join(paragraphs), len(paragraphs) // 30 + 1, "docx"
        except ImportError:
            logger.warning("python-docx not installed, falling back to stub extraction.")
            return self._stub_text(), 1, "docx"

    # ── Text Cleaning ─────────────────────────────────────────

    def _clean_text(self, text: str) -> str:
        text = re.sub(r"\r\n", "\n", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r"[ \t]{2,}", " ", text)
        text = re.sub(r"\x00", "", text)             # null bytes
        text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E]", " ", text)  # non-printable
        return text.strip()

    # ── Chunking ─────────────────────────────────────────────

    def _chunk_text(self, text: str) -> list[DocumentChunk]:
        chunks: list[DocumentChunk] = []
        start = 0
        idx = 0

        while start < len(text):
            end = min(start + CHUNK_SIZE, len(text))

            # Try to break at paragraph boundary
            if end < len(text):
                para_break = text.rfind("\n\n", start, end)
                if para_break > start + CHUNK_SIZE // 2:
                    end = para_break

            chunk_text = text[start:end].strip()
            if chunk_text:
                section_ref = self._detect_section(chunk_text)
                chunks.append(DocumentChunk(
                    index=idx,
                    text=chunk_text,
                    char_start=start,
                    char_end=end,
                    section_ref=section_ref,
                ))
                idx += 1

            start = end - CHUNK_OVERLAP if end < len(text) else len(text)

        return chunks

    def _detect_section(self, text: str) -> Optional[str]:
        m = _SECTION_RE.search(text)
        if m:
            return f"Section {m.group(1)}"
        return None

    # ── Stub for testing without actual files ─────────────────

    def _stub_text(self) -> str:
        return """NON-DISCLOSURE AGREEMENT

This Non-Disclosure Agreement ("Agreement") is entered into as of June 10, 2025.

SECTION 1. CONFIDENTIAL INFORMATION
The Receiving Party agrees to protect all Confidential Information.

SECTION 8. LIABILITY
Neither party's aggregate liability shall be UNLIMITED for any claims including indemnification.

SECTION 9. INDEMNIFICATION
Company shall indemnify and hold harmless the other party from ANY AND ALL claims with no limitation.

SECTION 10. TERMINATION
Either party may terminate upon 90 days written notice.

SECTION 11. GOVERNING LAW
This Agreement shall be governed by the laws of the State of Delaware."""
